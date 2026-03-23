"""
document_processor.py
=====================
Handles loading and chunking documents in multiple formats:
PDF, TXT, MD, CSV, JSON, DOCX.
"""

import csv
import hashlib
import io
import json
import os
import tempfile
from pathlib import Path

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS


def generate_doc_id(filename: str, content: str, index: int) -> str:
    """Content-addressed ID — same file + content always produces same ID."""
    raw = f"{filename}::{content[:500]}::{index}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


def load_file(uploaded_file) -> list[Document]:
    """
    Load an uploaded Streamlit file into LangChain Documents.
    Supports: PDF, TXT, MD, CSV, JSON, DOCX.
    """
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    # ── CSV ───────────────────────────────────────────────────────────
    if suffix == ".csv":
        return _load_csv(uploaded_file)

    # ── JSON ──────────────────────────────────────────────────────────
    if suffix == ".json":
        return _load_json(uploaded_file)

    # ── DOCX ──────────────────────────────────────────────────────────
    if suffix == ".docx":
        return _load_docx(uploaded_file)

    # ── PDF / TXT / MD — need temp file for LangChain loaders ────────
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name

    try:
        if suffix == ".pdf":
            loader = PyPDFLoader(tmp_path)
        elif suffix == ".md":
            loader = UnstructuredMarkdownLoader(tmp_path)
        elif suffix == ".txt":
            loader = TextLoader(tmp_path, encoding="utf-8")
        else:
            raise ValueError(f"Unhandled type: {suffix}")

        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = uploaded_file.name
        return docs

    finally:
        os.unlink(tmp_path)


def _load_csv(uploaded_file) -> list[Document]:
    """Load CSV — each row becomes a document with column headers as context."""
    content = uploaded_file.getvalue().decode("utf-8")
    reader = csv.DictReader(io.StringIO(content))
    docs = []

    for i, row in enumerate(reader):
        # Format each row as "Column: Value" pairs for better retrieval
        text = "\n".join(f"{k}: {v}" for k, v in row.items() if v)
        docs.append(Document(
            page_content=text,
            metadata={"source": uploaded_file.name, "row": i},
        ))

    return docs


def _load_json(uploaded_file) -> list[Document]:
    """Load JSON — handles objects, arrays of objects, and nested structures."""
    content = uploaded_file.getvalue().decode("utf-8")
    data = json.loads(content)

    docs = []
    if isinstance(data, list):
        for i, item in enumerate(data):
            text = json.dumps(item, indent=2) if isinstance(item, dict) else str(item)
            docs.append(Document(
                page_content=text,
                metadata={"source": uploaded_file.name, "index": i},
            ))
    elif isinstance(data, dict):
        # Flatten top-level keys into separate documents
        for key, value in data.items():
            text = f"{key}:\n{json.dumps(value, indent=2) if isinstance(value, (dict, list)) else str(value)}"
            docs.append(Document(
                page_content=text,
                metadata={"source": uploaded_file.name, "key": key},
            ))
    else:
        docs.append(Document(
            page_content=str(data),
            metadata={"source": uploaded_file.name},
        ))

    return docs


def _load_docx(uploaded_file) -> list[Document]:
    """Load DOCX using python-docx — extracts paragraphs and tables."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx support. "
            "Install with: pip install python-docx"
        )

    doc = DocxDocument(io.BytesIO(uploaded_file.getvalue()))
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row_data) if v)
            if row_text:
                parts.append(row_text)

    full_text = "\n\n".join(parts)
    return [Document(
        page_content=full_text,
        metadata={"source": uploaded_file.name},
    )]


def chunk_documents(docs: list[Document]) -> list[Document]:
    """Split documents into retrieval-sized chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)
